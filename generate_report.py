import os
import sys
from datetime import datetime
from collections import Counter
from docx import Document
from docx.shared import RGBColor, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

# Directly import the review function
from review_contract import review_contract

COLOR_MAP = {
    "GREEN": RGBColor(0x22, 0x8B, 0x22),
    "RED":   RGBColor(0xCC, 0x00, 0x00),
    "BLUE":  RGBColor(0x00, 0x5B, 0xB5),
    "AMBER": RGBColor(0xE6, 0x8A, 0x00),
}

STATUS_LABELS = {
    "GREEN": "COMPLIANT",
    "RED":   "VIOLATION — Manager alert required",
    "BLUE":  "NO MATCHING CLAUSE — Manager alert required",
    "AMBER": "UNCERTAIN — Historical check + Manager alert required",
}

MANAGER_STATUSES = {"RED", "BLUE", "AMBER"}

def generate_report(contract_path):
    # Run the review
    review_data = review_contract(contract_path)
    if not review_data:
        print("Review was cancelled, no report generated.")
        return

    contract_type = review_data["contract_type"]
    results = review_data["results"]

    # Generate report filename
    timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
    base_name = os.path.splitext(os.path.basename(contract_path))[0]
    output_path = f"reports/{base_name}_compliance_{timestamp}.docx"
    os.makedirs("reports", exist_ok=True)

    doc = Document()

    # ── Title Page ──────────────────────────────────────────────
    title = doc.add_heading("Contract Compliance Review Report", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph(f"Contract file   : {os.path.basename(contract_path)}")
    doc.add_paragraph(f"Contract type   : {contract_type}")
    doc.add_paragraph(f"Review date     : {datetime.now().strftime('%Y-%m-%d %H:%M')}")
    doc.add_paragraph(f"Reviewed by     : AI Contract Reviewer (gpt-4o-mini)")
    doc.add_paragraph("")

    # ── Executive Summary ────────────────────────────────────────────────
    doc.add_heading("Executive Summary", level=1)
    counts = Counter(r["status"] for r in results)
    total = len(results)

    for status in ["GREEN", "RED", "BLUE", "AMBER"]:
        count = counts.get(status, 0)
        p = doc.add_paragraph()
        run = p.add_run(f"  {status}: {count}/{total} clause(s)")
        run.font.color.rgb = COLOR_MAP[status]
        run.bold = True

    # Manager alert summary
    alerts = [r for r in results if r["status"] in MANAGER_STATUSES]
    if alerts:
        doc.add_paragraph("")
        alert_heading = doc.add_paragraph()
        alert_run = alert_heading.add_run(
            f"⚠️  {len(alerts)} clause(s) require manager intervention"
        )
        alert_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
        alert_run.bold = True

    doc.add_paragraph("")

    # ── Review Results ─────────────────────────────────────────
    doc.add_heading("Clause-by-Clause Review", level=1)

    for r in results:
        # Title for the clause (with color)
        p = doc.add_heading("", level=2)
        run = p.add_run(
            f"Clause {r['clause_number']}  |  {STATUS_LABELS[r['status']]}"
        )
        run.font.color.rgb = COLOR_MAP[r["status"]]
        run.font.size = Pt(12)

        # Manager alert
        if r["status"] in MANAGER_STATUSES:
            alert_p = doc.add_paragraph()
            a_run = alert_p.add_run("⚠️  Manager intervention required")
            a_run.font.color.rgb = RGBColor(0xCC, 0x00, 0x00)
            a_run.italic = True

        # Clause text
        doc.add_paragraph("Clause text:").bold = True
        doc.add_paragraph(r["clause_text"])

        # Review reasoning
        p2 = doc.add_paragraph()
        p2.add_run("Reasoning: ").bold = True
        p2.add_run(r["reasoning"])

        # Cited position clause
        p3 = doc.add_paragraph()
        p3.add_run("Reference cited:").bold = True
        cited = r["cited_position_clause"]
        if isinstance(cited, dict):
           doc.add_paragraph(f"  Document : {cited.get('document', 'N/A')}")
           doc.add_paragraph(f"  Section  : {cited.get('section', 'N/A')}")
           doc.add_paragraph(f"  Quote    : {cited.get('quote', 'None found')}")
        else:
           doc.add_paragraph(str(cited))
        

        # Historical precedent (AMBER only)
        if r.get("historical_precedent"):
            p4 = doc.add_paragraph()
            p4.add_run("Historical precedent: ").bold = True
            p4.add_run(r["historical_precedent"])

        doc.add_paragraph("─" * 60)

    doc.save(output_path)
    print(f"\n✅ Report saved: {output_path}")

if __name__ == "__main__":
    if len(sys.argv) < 2:
        print("Usage: python 3_generate_report.py <contract_file>")
    else:
        generate_report(sys.argv[1])